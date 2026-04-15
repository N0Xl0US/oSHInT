from docx import Document

doc = Document()
doc.add_paragraph("Alex Dev")
doc.add_paragraph("alex@example.com")
doc.add_paragraph("New York, NY")

# Add a text link
doc.add_paragraph("Check out my GitHub: https://github.com/alexdev123")
doc.add_paragraph("And my LinkedIn: https://linkedin.com/in/alexdev22")

doc.add_paragraph("EXPERIENCE")
doc.add_paragraph("Massachusetts Institute of Technology")
doc.save("test_resume.docx")
